"""
RAG (Retrieval-Augmented Generation) service for intelligent query processing.
"""
import asyncio
import hashlib
import json
from typing import List, Dict, Any, Optional, Tuple
from uuid import UUID, uuid4
from datetime import datetime

from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.vectorstores import FAISS
from langchain.embeddings import OpenAIEmbeddings
from langchain.chat_models import ChatOpenAI
from langchain.chains import ConversationalRetrievalChain
from langchain.memory import ConversationBufferWindowMemory
from langchain.prompts import PromptTemplate
from langchain.schema import Document as LangchainDocument

from app.core.config import settings
from app.core.exceptions import RAGException
from app.models.schemas import (
    ChatResponse, SearchResult, SearchFilters, Document,
    ChatMessage, ChatSession, DocumentStatus
)
from app.services.database import DatabaseService
from app.services.redis_client import RedisClient
from app.services.neo4j_client import Neo4jClient
from app.services.embedding_service import EmbeddingService
from app.utils.logger import get_logger

logger = get_logger(__name__)


class RAGService:
    """RAG service for intelligent query processing and response generation."""
    
    def __init__(
        self,
        db_service: DatabaseService,
        redis_client: RedisClient,
        neo4j_client: Neo4jClient,
        embedding_service: EmbeddingService,
    ):
        self.db_service = db_service
        self.redis_client = redis_client
        self.neo4j_client = neo4j_client
        self.embedding_service = embedding_service
        
        # Initialize LLM
        self.llm = ChatOpenAI(
            model_name=settings.OPENAI_MODEL,
            temperature=settings.OPENAI_TEMPERATURE,
            max_tokens=settings.OPENAI_MAX_TOKENS,
            openai_api_key=settings.OPENAI_API_KEY,
        )
        
        # Initialize text splitter
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=settings.CHUNK_SIZE,
            chunk_overlap=settings.CHUNK_OVERLAP,
            length_function=len,
        )
        
        # Initialize vector store
        self.vector_store = None
        self._initialize_vector_store()
    
    def _initialize_vector_store(self):
        """Initialize FAISS vector store."""
        try:
            # Try to load existing vector store
            self.vector_store = FAISS.load_local(
                "data/vectorstore/faiss_index",
                self.embedding_service.get_embeddings_model()
            )
            logger.info("Loaded existing vector store")
        except Exception as e:
            logger.warning(f"Could not load existing vector store: {e}")
            # Create new vector store
            self.vector_store = FAISS.from_texts(
                ["initialization"],
                self.embedding_service.get_embeddings_model()
            )
            logger.info("Created new vector store")
    
    async def process_message(
        self,
        message: str,
        session_id: UUID,
        user_id: Optional[str] = None,
        context: Optional[Dict[str, Any]] = None,
    ) -> ChatResponse:
        """
        Process a chat message and generate AI response.
        
        Args:
            message: User message
            session_id: Chat session ID
            user_id: User ID
            context: Additional context
            
        Returns:
            ChatResponse: AI response
        """
        try:
            start_time = datetime.utcnow()
            
            # Get or create session
            session = await self._get_or_create_session(session_id, user_id)
            
            # Get conversation history
            history = await self._get_conversation_history(session_id)
            
            # Enhanced query processing with knowledge graph
            enhanced_query = await self._enhance_query_with_kg(message, context)
            
            # Retrieve relevant documents
            retrieved_docs = await self._retrieve_relevant_documents(
                enhanced_query, history, limit=5
            )
            
            # Generate response using RAG
            response_content = await self._generate_response(
                enhanced_query, retrieved_docs, history
            )
            
            # Calculate processing time
            processing_time = (datetime.utcnow() - start_time).total_seconds()
            
            # Create response
            message_id = uuid4()
            response = ChatResponse(
                content=response_content,
                session_id=session_id,
                message_id=message_id,
                retrieved_documents=[doc.dict() for doc in retrieved_docs],
                confidence_score=self._calculate_confidence_score(retrieved_docs),
                processing_time=processing_time,
                metadata={
                    "enhanced_query": enhanced_query,
                    "context": context,
                    "timestamp": datetime.utcnow().isoformat(),
                }
            )
            
            return response
            
        except Exception as e:
            logger.error(f"Error processing message: {e}")
            raise RAGException(f"Failed to process message: {str(e)}")
    
    async def _get_or_create_session(self, session_id: UUID, user_id: Optional[str]) -> ChatSession:
        """Get or create chat session."""
        session = await self.db_service.get_chat_session(session_id)
        
        if not session:
            session = await self.db_service.create_chat_session(
                id=session_id,
                user_id=user_id,
                title=f"Chat Session {datetime.utcnow().strftime('%Y-%m-%d %H:%M')}",
                is_active=True,
                metadata={}
            )
        
        return session
    
    async def _get_conversation_history(self, session_id: UUID) -> List[ChatMessage]:
        """Get conversation history for session."""
        return await self.db_service.get_session_messages(session_id)
    
    async def _enhance_query_with_kg(
        self,
        query: str,
        context: Optional[Dict[str, Any]] = None
    ) -> str:
        """Enhance query using knowledge graph information."""
        try:
            # Extract entities from query
            entities = await self._extract_entities_from_query(query)
            
            # Get related entities from knowledge graph
            related_entities = []
            for entity in entities:
                kg_entity = await self.neo4j_client.get_node("Entity", "name", entity)
                if kg_entity:
                    related = await self.neo4j_client.get_node_relationships(
                        "Entity", kg_entity["id"], direction="both"
                    )
                    related_entities.extend(related)
            
            # Build enhanced query with context
            enhanced_query = query
            if related_entities:
                entity_context = self._build_entity_context(related_entities)
                enhanced_query = f"{query}\n\nContext: {entity_context}"
            
            return enhanced_query
            
        except Exception as e:
            logger.error(f"Error enhancing query with KG: {e}")
            return query
    
    async def _extract_entities_from_query(self, query: str) -> List[str]:
        """Extract entities from query using NLP."""
        try:
            # Use OpenAI to extract entities
            extraction_prompt = f"""
            Extract relevant entities (satellites, instruments, data products, locations, etc.) 
            from this MOSDAC-related query:
            
            Query: {query}
            
            Return entities as a comma-separated list:
            """
            
            response = await self.llm.apredict(extraction_prompt)
            entities = [e.strip() for e in response.split(",") if e.strip()]
            
            return entities
            
        except Exception as e:
            logger.error(f"Error extracting entities: {e}")
            return []
    
    def _build_entity_context(self, related_entities: List[Dict[str, Any]]) -> str:
        """Build context string from related entities."""
        context_parts = []
        
        for entity_data in related_entities[:5]:  # Limit context
            entity = entity_data.get("related", {})
            relationship = entity_data.get("r", {})
            
            if entity and relationship:
                context_parts.append(
                    f"{entity.get('name', 'Unknown')} "
                    f"({entity.get('type', 'entity')}) "
                    f"- {relationship.get('type', 'related')}"
                )
        
        return "; ".join(context_parts)
    
    async def _retrieve_relevant_documents(
        self,
        query: str,
        history: List[ChatMessage],
        limit: int = 5
    ) -> List[SearchResult]:
        """Retrieve relevant documents for query."""
        try:
            # Check cache first
            cache_key = f"search_results:{hashlib.md5(query.encode()).hexdigest()}"
            cached_results = await self.redis_client.get(cache_key)
            
            if cached_results:
                return [SearchResult(**result) for result in cached_results]
            
            # Perform semantic search
            semantic_results = await self._semantic_search(query, limit)
            
            # Perform keyword search
            keyword_results = await self._keyword_search(query, limit)
            
            # Combine and rank results
            combined_results = self._combine_search_results(
                semantic_results, keyword_results, limit
            )
            
            # Cache results
            await self.redis_client.set(
                cache_key,
                [result.dict() for result in combined_results],
                ttl=300
            )
            
            return combined_results
            
        except Exception as e:
            logger.error(f"Error retrieving documents: {e}")
            return []
    
    async def _semantic_search(self, query: str, limit: int) -> List[SearchResult]:
        """Perform semantic search using embeddings."""
        try:
            if not self.vector_store:
                return []
            
            # Perform similarity search
            docs = self.vector_store.similarity_search_with_score(query, k=limit)
            
            results = []
            for doc, score in docs:
                result = SearchResult(
                    id=uuid4(),
                    title=doc.metadata.get("title", ""),
                    content=doc.page_content,
                    relevance_score=1.0 - score,  # Convert distance to relevance
                    document_type=doc.metadata.get("document_type", "text"),
                    category=doc.metadata.get("category"),
                    tags=doc.metadata.get("tags", []),
                    created_at=datetime.utcnow(),
                    source_url=doc.metadata.get("source_url"),
                    metadata=doc.metadata,
                )
                results.append(result)
            
            return results
            
        except Exception as e:
            logger.error(f"Error in semantic search: {e}")
            return []
    
    async def _keyword_search(self, query: str, limit: int) -> List[SearchResult]:
        """Perform keyword search in database."""
        try:
            documents = await self.db_service.search_documents(
                query=query,
                limit=limit,
                filters={"status": DocumentStatus.COMPLETED}
            )
            
            results = []
            for doc in documents:
                result = SearchResult(
                    id=doc.id,
                    title=doc.title,
                    content=doc.content[:500] + "..." if len(doc.content) > 500 else doc.content,
                    relevance_score=0.8,  # Default relevance for keyword search
                    document_type=doc.content_type,
                    category=doc.category,
                    tags=doc.tags,
                    created_at=doc.created_at,
                    source_url=doc.source_url,
                    metadata={
                        "document_id": str(doc.id),
                        "file_path": doc.file_path,
                        "processed_at": doc.processed_at.isoformat() if doc.processed_at else None,
                    },
                )
                results.append(result)
            
            return results
            
        except Exception as e:
            logger.error(f"Error in keyword search: {e}")
            return []
    
    def _combine_search_results(
        self,
        semantic_results: List[SearchResult],
        keyword_results: List[SearchResult],
        limit: int
    ) -> List[SearchResult]:
        """Combine and rank search results."""
        # Create a map to avoid duplicates
        results_map = {}
        
        # Add semantic results with higher weight
        for result in semantic_results:
            key = f"{result.title}:{result.content[:100]}"
            if key not in results_map:
                result.relevance_score *= 1.2  # Boost semantic results
                results_map[key] = result
        
        # Add keyword results
        for result in keyword_results:
            key = f"{result.title}:{result.content[:100]}"
            if key not in results_map:
                results_map[key] = result
        
        # Sort by relevance score and return top results
        combined_results = sorted(
            results_map.values(),
            key=lambda x: x.relevance_score,
            reverse=True
        )
        
        return combined_results[:limit]
    
    async def _generate_response(
        self,
        query: str,
        retrieved_docs: List[SearchResult],
        history: List[ChatMessage]
    ) -> str:
        """Generate response using RAG."""
        try:
            # Build context from retrieved documents
            context = self._build_context_from_documents(retrieved_docs)
            
            # Build conversation history
            conversation_history = self._build_conversation_history(history)
            
            # Create prompt
            prompt = self._create_rag_prompt(query, context, conversation_history)
            
            # Generate response
            response = await self.llm.apredict(prompt)
            
            return response.strip()
            
        except Exception as e:
            logger.error(f"Error generating response: {e}")
            return "I apologize, but I encountered an error while processing your query. Please try again."
    
    def _build_context_from_documents(self, documents: List[SearchResult]) -> str:
        """Build context string from retrieved documents."""
        context_parts = []
        
        for i, doc in enumerate(documents, 1):
            context_parts.append(
                f"[Source {i}] {doc.title}\n"
                f"Content: {doc.content}\n"
                f"Category: {doc.category or 'General'}\n"
                f"---"
            )
        
        return "\n".join(context_parts)
    
    def _build_conversation_history(self, history: List[ChatMessage]) -> str:
        """Build conversation history string."""
        if not history:
            return ""
        
        history_parts = []
        for msg in history[-6:]:  # Last 6 messages
            role = "User" if msg.message_type == "user" else "Assistant"
            history_parts.append(f"{role}: {msg.content}")
        
        return "\n".join(history_parts)
    
    def _create_rag_prompt(
        self,
        query: str,
        context: str,
        conversation_history: str
    ) -> str:
        """Create RAG prompt for response generation."""
        prompt = f"""
You are MOSDAC AI Assistant, an expert on India's Meteorological and Oceanographic Satellite Data.
Your role is to provide accurate, helpful responses about satellite data, weather information, and MOSDAC services.

CONTEXT FROM DOCUMENTS:
{context}

CONVERSATION HISTORY:
{conversation_history}

CURRENT QUERY: {query}

INSTRUCTIONS:
1. Use the provided context to answer the query accurately
2. If the context doesn't contain enough information, clearly state what information is missing
3. Maintain conversation continuity by referencing previous messages when relevant
4. Provide specific, actionable information when possible
5. If discussing satellite data, include relevant technical details and usage guidelines
6. Format your response clearly with proper structure
7. Always be helpful and professional

RESPONSE:
"""
        return prompt
    
    def _calculate_confidence_score(self, retrieved_docs: List[SearchResult]) -> float:
        """Calculate confidence score based on retrieved documents."""
        if not retrieved_docs:
            return 0.0
        
        # Calculate average relevance score
        avg_relevance = sum(doc.relevance_score for doc in retrieved_docs) / len(retrieved_docs)
        
        # Factor in number of documents
        doc_factor = min(len(retrieved_docs) / 5.0, 1.0)
        
        # Combine factors
        confidence = (avg_relevance * 0.7) + (doc_factor * 0.3)
        
        return min(confidence, 1.0)
    
    # Document processing methods
    async def upload_document(
        self,
        file: Any,
        metadata: Dict[str, Any],
        background_tasks: Any,
    ) -> str:
        """Upload and process document."""
        try:
            # Generate document ID
            document_id = str(uuid4())
            
            # Read file content
            content = await file.read()
            
            # Create document record
            document = await self.db_service.create_document(
                id=document_id,
                title=metadata.get("title", file.filename),
                content_type=file.content_type,
                file_size=len(content),
                status=DocumentStatus.PROCESSING,
                description=metadata.get("description"),
                tags=metadata.get("tags", []),
                category=metadata.get("category"),
                language=metadata.get("language", "en"),
                source_type="uploaded",
                uploaded_by_id=metadata.get("uploaded_by"),
            )
            
            # Process document in background
            background_tasks.add_task(
                self._process_document_background,
                document_id,
                content,
                file.content_type
            )
            
            return document_id
            
        except Exception as e:
            logger.error(f"Error uploading document: {e}")
            raise RAGException(f"Failed to upload document: {str(e)}")
    
    async def _process_document_background(
        self,
        document_id: str,
        content: bytes,
        content_type: str
    ):
        """Process document in background."""
        try:
            # Extract text content
            text_content = await self._extract_text_content(content, content_type)
            
            # Update document with content
            await self.db_service.update(
                Document,
                document_id,
                content=text_content,
                processed_at=datetime.utcnow()
            )
            
            # Create document chunks
            chunks = self.text_splitter.split_text(text_content)
            
            # Store chunks in database
            for i, chunk in enumerate(chunks):
                chunk_hash = hashlib.md5(chunk.encode()).hexdigest()
                
                await self.db_service.create_document_chunk(
                    document_id=document_id,
                    chunk_index=i,
                    content=chunk,
                    content_hash=chunk_hash,
                    embedding_model=settings.EMBEDDING_MODEL,
                    embedding_dimension=settings.EMBEDDING_DIMENSION,
                    token_count=len(chunk.split()),
                )
            
            # Generate embeddings and update vector store
            await self._update_vector_store_with_document(document_id, chunks)
            
            # Update document status
            await self.db_service.update(
                Document,
                document_id,
                status=DocumentStatus.COMPLETED
            )
            
            logger.info(f"Document {document_id} processed successfully")
            
        except Exception as e:
            logger.error(f"Error processing document {document_id}: {e}")
            
            # Update document status to failed
            await self.db_service.update(
                Document,
                document_id,
                status=DocumentStatus.FAILED
            )
    
    async def _extract_text_content(self, content: bytes, content_type: str) -> str:
        """Extract text content from various file types."""
        try:
            if content_type == "text/plain":
                return content.decode("utf-8")
            elif content_type == "application/pdf":
                # Use PyPDF2 or similar library
                import PyPDF2
                import io
                
                pdf_file = io.BytesIO(content)
                pdf_reader = PyPDF2.PdfReader(pdf_file)
                
                text = ""
                for page in pdf_reader.pages:
                    text += page.extract_text()
                
                return text
            elif content_type in ["application/msword", "application/vnd.openxmlformats-officedocument.wordprocessingml.document"]:
                # Use python-docx
                from docx import Document
                import io
                
                doc_file = io.BytesIO(content)
                doc = Document(doc_file)
                
                text = ""
                for paragraph in doc.paragraphs:
                    text += paragraph.text + "\n"
                
                return text
            else:
                # Try to decode as text
                return content.decode("utf-8", errors="ignore")
                
        except Exception as e:
            logger.error(f"Error extracting text content: {e}")
            return ""
    
    async def _update_vector_store_with_document(
        self,
        document_id: str,
        chunks: List[str]
    ):
        """Update vector store with document chunks."""
        try:
            # Get document metadata
            document = await self.db_service.get_document(document_id)
            
            if not document:
                return
            
            # Create Langchain documents
            langchain_docs = []
            for i, chunk in enumerate(chunks):
                metadata = {
                    "document_id": document_id,
                    "title": document.title,
                    "chunk_index": i,
                    "document_type": document.content_type,
                    "category": document.category,
                    "tags": document.tags,
                    "source_url": document.source_url,
                }
                
                langchain_docs.append(
                    LangchainDocument(page_content=chunk, metadata=metadata)
                )
            
            # Add to vector store
            if langchain_docs:
                self.vector_store.add_documents(langchain_docs)
                
                # Save updated vector store
                self.vector_store.save_local("data/vectorstore/faiss_index")
                
                logger.info(f"Added {len(langchain_docs)} chunks to vector store")
                
        except Exception as e:
            logger.error(f"Error updating vector store: {e}")
    
    # Session management
    async def get_user_sessions(self, user_id: str) -> List[ChatSession]:
        """Get user's chat sessions."""
        return await self.db_service.get_user_sessions(user_id)
    
    async def get_session_messages(self, session_id: UUID, user_id: str) -> List[ChatMessage]:
        """Get messages for a session."""
        # Verify session belongs to user
        session = await self.db_service.get_chat_session(session_id)
        if not session or session.user_id != user_id:
            raise RAGException("Session not found or access denied")
        
        return await self.db_service.get_session_messages(session_id)
    
    async def delete_session(self, session_id: UUID, user_id: str) -> bool:
        """Delete a chat session."""
        # Verify session belongs to user
        session = await self.db_service.get_chat_session(session_id)
        if not session or session.user_id != user_id:
            raise RAGException("Session not found or access denied")
        
        return await self.db_service.delete(ChatSession, session_id)
    
    async def store_conversation(
        self,
        session_id: UUID,
        user_message: str,
        ai_response: str,
        user_id: Optional[str] = None
    ):
        """Store conversation in database."""
        try:
            # Store user message
            await self.db_service.create_chat_message(
                session_id=session_id,
                message_type="user",
                content=user_message,
                metadata={"timestamp": datetime.utcnow().isoformat()}
            )
            
            # Store AI response
            await self.db_service.create_chat_message(
                session_id=session_id,
                message_type="assistant",
                content=ai_response,
                metadata={"timestamp": datetime.utcnow().isoformat()}
            )
            
        except Exception as e:
            logger.error(f"Error storing conversation: {e}")
    
    async def store_feedback(
        self,
        message_id: UUID,
        rating: int,
        feedback: Optional[str] = None,
        user_id: Optional[str] = None
    ):
        """Store user feedback."""
        try:
            await self.db_service.create_feedback(
                message_id=message_id,
                user_id=user_id,
                rating=rating,
                feedback_text=feedback,
                feedback_type="rating"
            )
            
        except Exception as e:
            logger.error(f"Error storing feedback: {e}")
    
    # Search methods
    async def search_documents(
        self,
        query: str,
        search_type: str = "hybrid",
        filters: Optional[SearchFilters] = None,
        limit: int = 10,
        offset: int = 0,
        include_metadata: bool = True,
        user_id: Optional[str] = None,
    ) -> Tuple[List[SearchResult], int]:
        """Search documents with various strategies."""
        try:
            # Log search query
            await self.db_service.create_search_query(
                query=query,
                search_type=search_type,
                results_count=0,  # Will be updated
                execution_time=0.0,  # Will be updated
                user_id=user_id,
            )
            
            if search_type == "semantic":
                results = await self._semantic_search(query, limit + offset)
                results = results[offset:offset + limit]
                total = len(results)
            elif search_type == "keyword":
                results = await self._keyword_search(query, limit + offset)
                results = results[offset:offset + limit]
                total = len(results)
            else:  # hybrid
                semantic_results = await self._semantic_search(query, limit)
                keyword_results = await self._keyword_search(query, limit)
                results = self._combine_search_results(semantic_results, keyword_results, limit)
                total = len(results)
            
            return results, total
            
        except Exception as e:
            logger.error(f"Error searching documents: {e}")
            return [], 0
    
    async def get_search_suggestions(self, query_prefix: str, limit: int = 10) -> List[str]:
        """Get search suggestions based on query prefix."""
        try:
            # Get trending queries that match prefix
            trending = await self.db_service.get_trending_queries(limit * 2)
            
            # Filter by prefix
            suggestions = [
                q for q in trending
                if q.lower().startswith(query_prefix.lower())
            ]
            
            return suggestions[:limit]
            
        except Exception as e:
            logger.error(f"Error getting search suggestions: {e}")
            return []
    
    async def get_available_filters(self) -> Dict[str, Any]:
        """Get available search filters."""
        try:
            # This would typically query the database for available values
            filters = {
                "categories": ["Weather", "Ocean", "Satellite", "Climate", "General"],
                "content_types": ["text/plain", "application/pdf", "application/json"],
                "languages": ["en", "hi"],
                "source_types": ["scraped", "uploaded", "api"],
                "date_ranges": {
                    "min_date": "2020-01-01",
                    "max_date": datetime.utcnow().strftime("%Y-%m-%d")
                }
            }
            
            return filters
            
        except Exception as e:
            logger.error(f"Error getting available filters: {e}")
            return {}
    
    async def find_similar_documents(
        self,
        document_id: str,
        limit: int = 10
    ) -> List[SearchResult]:
        """Find documents similar to a given document."""
        try:
            # Get document content
            document = await self.db_service.get_document(document_id)
            if not document:
                return []
            
            # Use document title and content for similarity search
            query = f"{document.title} {document.content[:500]}"
            
            # Perform semantic search
            results = await self._semantic_search(query, limit + 1)
            
            # Remove the original document from results
            filtered_results = [
                r for r in results
                if r.metadata.get("document_id") != document_id
            ]
            
            return filtered_results[:limit]
            
        except Exception as e:
            logger.error(f"Error finding similar documents: {e}")
            return []
    
    async def get_trending_queries(self, limit: int = 10) -> List[str]:
        """Get trending search queries."""
        return await self.db_service.get_trending_queries(limit)
